
import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import unicodedata

COMMERCIAUX_CIBLES = ['Sandra', 'Ophélie', 'Arthur', 'Grégoire', 'Tania']
PARTIES = [
    ('Sites créés', 'sites_crees', True),
    ('Offres à remettre', 'offres_a_remettre_detail', False),
    ('PDB à remettre', 'pdbs_a_remettre_detail', False),
    ('Offres signées', 'offre_signee_detail', True),
    ('PDB signées', 'pdbs_signees', True)
]

def sanitize_filename(name):
    return name.replace(" ", "_").replace("/", "-")

def normalize(text):
    return ''.join(c for c in unicodedata.normalize('NFD', str(text)) if unicodedata.category(c) != 'Mn').lower().replace('_', ' ').replace('-', ' ')

def detect_column(columns, keyword):
    keyword_norm = normalize(keyword)
    for col in columns:
        if keyword_norm in normalize(col):
            return col
    return None

def convert_mois_to_int(val):
    if pd.isnull(val):
        return None
    val = str(val).strip().lower()
    mois_dict = {
        'janvier': 1, 'février': 2, 'mars': 3, 'avril': 4,
        'mai': 5, 'juin': 6, 'juillet': 7, 'août': 8,
        'septembre': 9, 'octobre': 10, 'novembre': 11, 'décembre': 12,
        'january': 1, 'february': 2, 'march': 3, 'april': 4,
        'may': 5, 'june': 6, 'july': 7, 'august': 8,
        'september': 9, 'october': 10, 'november': 11, 'december': 12
    }
    return mois_dict.get(val, None)

def charger_donnees(excel_path, mois_cible, annee_cible, jour_debut=None, jour_fin=None):
    xls = pd.ExcelFile(excel_path)
    data_by_part = {}

    for titre, sheet, _ in PARTIES:
        try:
            df = pd.read_excel(xls, sheet_name=sheet)
        except:
            continue

        col_annee = detect_column(df.columns, 'annee')
        col_mois = detect_column(df.columns, 'mois')
        col_com = detect_column(df.columns, 'commercial')

        if not col_annee or not col_mois or not col_com:
            continue

        df[col_mois] = df[col_mois].apply(convert_mois_to_int)
        df_filtre = df[(df[col_annee] == annee_cible) & (df[col_mois] == mois_cible)]
        col_jour = detect_column(df.columns, 'jour')
        if col_jour and jour_debut and jour_fin:
            df_filtre = df_filtre[(df_filtre[col_jour] >= jour_debut) & (df_filtre[col_jour] <= jour_fin)]
        df_filtre = df_filtre[df_filtre[col_com].str.contains('|'.join(COMMERCIAUX_CIBLES), case=False, na=False)]
        if df_filtre.empty:
            continue

        data_by_part[titre] = dict(tuple(df_filtre.groupby(col_com)))
    return data_by_part

def ajouter_page_de_garde(doc, nom_commercial, jour_debut, jour_fin, mois, annee, logo_path=None):
    doc.add_paragraph()
    if logo_path and os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(2))
    doc.add_heading("RAPPORT COMMERCIAL", level=0)
    doc.add_paragraph(f"Entreprise : Watt&Co Ingénierie")
    doc.add_paragraph(f"Commercial : {nom_commercial}")
    doc.add_paragraph(f"Période : du {jour_debut} au {jour_fin} {datetime(annee, mois, 1).strftime('%B')} {annee}")
    doc.add_paragraph("Contenu du rapport :")
    for titre, _, _ in PARTIES:
        doc.add_paragraph(f"• {titre}", style="List Bullet")
    doc.add_page_break()


def ajouter_logo_et_titre(doc, logo_path, nom, jour_debut, jour_fin, mois, annee):
    header = doc.sections[0].header
    p = header.paragraphs[0]
    if logo_path and os.path.exists(logo_path):
        p.add_run().add_picture(logo_path, width=Inches(1.5))
    mois_nom = datetime(annee, mois, 1).strftime('%B')
    p.add_run(f"   Compte rendu du {jour_debut} au {jour_fin} {mois_nom} {annee} – Réunion commerciale {nom}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def ajouter_statistiques_mensuelles(doc, titre, df, mois, annee):
    para = doc.add_paragraph()
    para.add_run(f"Année : {annee}\n").bold = True
    para.add_run(f"Mois : {mois}\n").bold = True
    para.add_run(f"Nombre de {titre.lower()} : {len(df)}\n").bold = True
    col_puissance = detect_column(df.columns, 'puissance')
    if col_puissance:
        total_puissance = df[col_puissance].sum()
        para.add_run(f"Puissance totale : {total_puissance:.2f} kWc").bold = True

def ajouter_tableau(doc, df, exclure=[]):
    table = doc.add_table(rows=1, cols=len(df.columns) - len(exclure))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    headers = [col for col in df.columns if col not in exclure]

    for i, col in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = col
        cell.paragraphs[0].runs[0].font.size = Pt(7)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade = OxmlElement('w:shd')
        shade.set(qn('w:fill'), 'D9E1F2')
        cell._tc.get_or_add_tcPr().append(shade)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(headers):
            row_cells[i].text = str(row[col]) if pd.notnull(row[col]) else ''
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def creer_graphique_global(excel_path, sheet, commercial, img_path):
    df = pd.read_excel(excel_path, sheet_name=sheet)
    col_mois = detect_column(df.columns, 'mois')
    col_annee = detect_column(df.columns, 'annee')
    col_com = detect_column(df.columns, 'commercial')

    if not col_mois or not col_annee or not col_com:
        return

    df[col_mois] = df[col_mois].apply(convert_mois_to_int)
    df_filtre = df[(df[col_annee] == datetime.now().year) & (df[col_com].str.contains(commercial, case=False, na=False))]
    if df_filtre.empty:
        return

    counts = df_filtre.groupby(df_filtre[col_mois]).size().reindex(range(1, 13), fill_value=0)
    months = ['Jan', 'Fév', 'Mar', 'Avr', 'Mai', 'Juin', 'Juil', 'Août', 'Sep', 'Oct', 'Nov', 'Déc']
    plt.figure(figsize=(6, 3))
    plt.bar(months, counts.values, color="#4F81BD")
    plt.title(f"Évolution mensuelle – {commercial}")
    plt.xlabel('Mois')
    plt.ylabel('Nombre')
    plt.tight_layout()
    plt.savefig(img_path)
    plt.close()

def plot_puissance(excel_path, sheet_name, commercial, output_path):
    df = pd.read_excel(excel_path, sheet_name=sheet_name)
    col_mois = detect_column(df.columns, 'mois')
    col_annee = detect_column(df.columns, 'annee')
    col_com = detect_column(df.columns, 'commercial')
    col_puissance = detect_column(df.columns, 'puissance')

    if not col_mois or not col_annee or not col_com or not col_puissance:
        return

    df[col_mois] = df[col_mois].apply(convert_mois_to_int)
    df_filtre = df[(df[col_annee] == datetime.now().year) & (df[col_com].str.contains(commercial, case=False, na=False))]
    if df_filtre.empty:
        return

    puissances = df_filtre.groupby(df_filtre[col_mois])[col_puissance].sum().reindex(range(1, 13), fill_value=0)
    mois_labels = ['Jan', 'Fév', 'Mar', 'Avr', 'Mai', 'Juin', 'Juil', 'Août', 'Sep', 'Oct', 'Nov', 'Déc']
    plt.figure(figsize=(6, 3))
    plt.plot(mois_labels, puissances.values, marker='o', color='green', linewidth=2)
    plt.title(f"Puissance mensuelle – {commercial}")
    plt.xlabel('Mois')
    plt.ylabel('Puissance (kWc)')
    plt.tight_layout()
    plt.savefig(output_path)
    plt.close()

def ajouter_section(doc, excel_path, titre, df, graphique, commercial, mois, annee, jour_debut, jour_fin, img_dir):
    doc.add_page_break()
    mois_nom = datetime(annee, mois, 1).strftime('%B')
    titre_complet = f"{titre} du {jour_debut} au {jour_fin} {mois_nom} {annee}"
    doc.add_heading(titre_complet, level=2)
    ajouter_statistiques_mensuelles(doc, titre, df, mois, annee)
    ajouter_tableau(doc, df, exclure=['lien'])
    doc.add_paragraph()
    if graphique:
        sheet = next((s for t, s, g in PARTIES if t == titre), None)
        if sheet:
            img_nb = os.path.join(img_dir, f"{sanitize_filename(commercial)}_{sanitize_filename(titre)}.png")
            creer_graphique_global(excel_path, sheet, commercial, img_nb)
            if os.path.exists(img_nb):
                doc.add_picture(img_nb, width=Inches(5))
                os.remove(img_nb)
            img_p = os.path.join(img_dir, f"{sanitize_filename(commercial)}_{sanitize_filename(titre)}_puissance.png")
            plot_puissance(excel_path, sheet, commercial, img_p)
            if os.path.exists(img_p):
                doc.add_picture(img_p, width=Inches(5))
                os.remove(img_p)

def creer_rapport(commercial, data_by_part, mois, annee, jour_debut, jour_fin, output_dir, excel_path, logo_path, img_dir):
    doc = Document()
    ajouter_page_de_garde(doc, commercial, jour_debut, jour_fin, mois, annee, logo_path)
    ajouter_logo_et_titre(doc, logo_path, commercial, datetime(annee, mois, 1), jour_debut, jour_fin)
    for titre, _, graphique in PARTIES:
        if commercial in data_by_part.get(titre, {}):
            titre_complet = f"{titre} du {jour_debut} au {jour_fin} {datetime(annee, mois, 1).strftime('%B')} {annee}"
            ajouter_section(doc, excel_path, titre_complet, data_by_part[titre][commercial], graphique, commercial, mois, annee, jour_debut, jour_fin, img_dir)
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/Rapport_Commercial_{sanitize_filename(commercial)}_{mois:02d}_{annee}.docx"
    doc.save(filename)
